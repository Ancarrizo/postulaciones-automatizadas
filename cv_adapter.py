"""
Adapta el CV .docx para cada postulación:
- Inserta un resumen profesional personalizado al principio
- Resalta keywords de la oferta en la sección de habilidades

El CV original nunca se modifica — se genera una copia temporal por postulación.
"""

import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

CVS_DIR  = Path(__file__).parent / "cvs"
TEMP_DIR = Path(__file__).parent / "cvs" / "temp"


def adaptar_cv(cv_filename: str, resumen: str, keywords: list[str], empresa: str) -> str:
    """
    Crea una copia del CV con el resumen adaptado insertado al principio.
    Retorna el nombre del archivo temporal generado.
    Solo funciona con .docx — los .pdf se usan tal cual.
    """
    cv_path = CVS_DIR / cv_filename

    # Si es PDF o no existe, devolver el original sin modificar
    if not cv_path.exists() or not cv_filename.endswith(".docx"):
        return cv_filename

    TEMP_DIR.mkdir(exist_ok=True)

    empresa_slug = empresa.replace(" ", "_").replace("/", "-")[:30]
    temp_name    = f"CV_Andres_{empresa_slug}.docx"
    temp_path    = TEMP_DIR / temp_name

    # Copiar el original
    shutil.copy2(cv_path, temp_path)

    doc = Document(str(temp_path))

    # ── Insertar resumen adaptado al principio del documento ─────────────────
    # Buscar el primer párrafo con contenido real (saltear encabezados vacíos)
    insert_idx = 0
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            insert_idx = i
            break

    # Insertar 2 párrafos antes del contenido: título + resumen
    _insert_paragraph_before(doc, insert_idx, "PERFIL PROFESIONAL", bold=True, size=11, space_after=2)
    _insert_paragraph_before(doc, insert_idx + 1, resumen, bold=False, size=10, space_after=6)
    _insert_paragraph_before(doc, insert_idx + 2, "", bold=False, size=6, space_after=0)  # separador

    doc.save(str(temp_path))
    return f"temp/{temp_name}"


def _insert_paragraph_before(doc: Document, idx: int, text: str,
                               bold: bool, size: int, space_after: int) -> None:
    """Inserta un párrafo en la posición idx del documento."""
    from docx.oxml.ns import qn
    from docx.oxml   import OxmlElement
    import copy

    # Crear nuevo párrafo
    new_para = OxmlElement("w:p")
    new_run  = OxmlElement("w:r")
    new_rpr  = OxmlElement("w:rPr")

    if bold:
        bold_el = OxmlElement("w:b")
        new_rpr.append(bold_el)

    sz_el = OxmlElement("w:sz")
    sz_el.set(qn("w:val"), str(size * 2))
    new_rpr.append(sz_el)

    new_run.append(new_rpr)

    new_t = OxmlElement("w:t")
    new_t.text = text
    new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    new_run.append(new_t)
    new_para.append(new_run)

    # Insertar antes del párrafo en idx
    ref_para = doc.paragraphs[min(idx, len(doc.paragraphs) - 1)]._element
    ref_para.getparent().insert(list(ref_para.getparent()).index(ref_para), new_para)


def limpiar_temporales() -> None:
    """Borra los CVs temporales generados."""
    if TEMP_DIR.exists():
        shutil.rmtree(TEMP_DIR)
        TEMP_DIR.mkdir()
